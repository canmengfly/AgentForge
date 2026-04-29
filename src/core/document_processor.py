from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import BinaryIO


@dataclass
class DocumentChunk:
    chunk_id: str
    doc_id: str
    content: str
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    doc_id: str
    title: str
    source: str
    content: str
    metadata: dict = field(default_factory=dict)


def _make_id(text: str) -> str:
    return hashlib.sha256(text.encode()).hexdigest()[:16]


def parse_text(content: str, title: str, source: str = "inline") -> ParsedDocument:
    doc_id = _make_id(title + content[:100])
    return ParsedDocument(doc_id=doc_id, title=title, source=source, content=content)


def parse_pdf(file: BinaryIO, filename: str) -> ParsedDocument:
    try:
        import pdfplumber
    except ImportError as e:
        raise RuntimeError("pdfplumber required: pip install pdfplumber") from e

    pages = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
    content = "\n\n".join(pages)
    doc_id = _make_id(filename + content[:100])
    return ParsedDocument(
        doc_id=doc_id,
        title=Path(filename).stem,
        source=filename,
        content=content,
        metadata={"page_count": len(pages)},
    )


def parse_docx(file: BinaryIO, filename: str) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx required: pip install python-docx") from e

    doc = Document(file)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(paragraphs)
    doc_id = _make_id(filename + content[:100])
    return ParsedDocument(
        doc_id=doc_id,
        title=Path(filename).stem,
        source=filename,
        content=content,
    )


def parse_html(content: str | bytes, filename: str) -> ParsedDocument:
    try:
        from bs4 import BeautifulSoup
    except ImportError as e:
        raise RuntimeError("beautifulsoup4 required: pip install beautifulsoup4") from e

    soup = BeautifulSoup(content, "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    title_tag = soup.find("title")
    title = title_tag.get_text().strip() if title_tag else Path(filename).stem
    doc_id = _make_id(filename + text[:100])
    return ParsedDocument(doc_id=doc_id, title=title, source=filename, content=text)


def parse_file(file: BinaryIO, filename: str) -> ParsedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(file, filename)
    if suffix == ".docx":
        return parse_docx(file, filename)
    if suffix in (".html", ".htm"):
        return parse_html(file.read(), filename)
    # Plain text / markdown
    raw = file.read()
    content = raw.decode("utf-8", errors="replace")
    doc_id = _make_id(filename + content[:100])
    return ParsedDocument(
        doc_id=doc_id,
        title=Path(filename).stem,
        source=filename,
        content=content,
    )


_SENT_END_RE = re.compile(r"(?<=[。！？!?…])\s*|(?<=\n)\n+")


def _split_sentences(text: str) -> list[str]:
    """Split at sentence / paragraph boundaries; preserve non-empty segments."""
    parts = _SENT_END_RE.split(text)
    return [p.strip() for p in parts if p.strip()]


def chunk_document(doc: ParsedDocument, chunk_size: int = 512, overlap: int = 64) -> list[DocumentChunk]:
    text = doc.content.strip()
    if not text:
        return []

    sentences = _split_sentences(text) or [text]

    # Greedily pack sentences into raw chunks
    raw_chunks: list[str] = []
    buf: list[str] = []
    buf_len = 0

    for sent in sentences:
        if len(sent) > chunk_size:
            # Flush buffer, then force-split the oversized sentence by characters
            if buf:
                raw_chunks.append(" ".join(buf))
                buf, buf_len = [], 0
            for i in range(0, len(sent), chunk_size):
                raw_chunks.append(sent[i: i + chunk_size])
            continue

        join_cost = 1 if buf else 0
        if buf_len + join_cost + len(sent) <= chunk_size:
            buf.append(sent)
            buf_len += join_cost + len(sent)
        else:
            if buf:
                raw_chunks.append(" ".join(buf))
            buf, buf_len = [sent], len(sent)

    if buf:
        raw_chunks.append(" ".join(buf))

    if not raw_chunks:
        raw_chunks = [text]

    total = len(raw_chunks)
    chunks: list[DocumentChunk] = []

    for i, raw in enumerate(raw_chunks):
        if i > 0 and overlap > 0:
            prev = raw_chunks[i - 1]
            tail = prev[-overlap:]
            # Trim to first whitespace so overlap starts at a word boundary
            ws = tail.find(" ")
            if 0 < ws < len(tail) - 1:
                tail = tail[ws + 1:]
            chunk_text = (tail + " " + raw).strip() if tail.strip() else raw
        else:
            chunk_text = raw

        chunk_id = _make_id(doc.doc_id + str(i))
        chunks.append(
            DocumentChunk(
                chunk_id=chunk_id,
                doc_id=doc.doc_id,
                content=chunk_text,
                metadata={
                    **doc.metadata,
                    "title": doc.title,
                    "source": doc.source,
                    "chunk_index": i,
                    "total_chunks": total,
                },
            )
        )

    return chunks
